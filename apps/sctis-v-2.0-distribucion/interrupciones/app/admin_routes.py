from flask import Blueprint, render_template, request, jsonify, g, send_from_directory, send_file, session
from app.db import query, query_one
from app.auth import login_required, role_required
from werkzeug.security import generate_password_hash
from datetime import datetime, date
import os
import io
import csv
import json
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'uploads')

bp = Blueprint('admin', __name__)


@bp.route('/admin/usuarios')
@login_required
@role_required('admin')
def usuarios_page():
    estados = query("SELECT state_code, state_name FROM common.states ORDER BY state_name")
    roles = query("SELECT * FROM sctis.user_roles ORDER BY role_name")
    return render_template('admin/usuarios.html', estados=estados, roles=roles)


@bp.route('/api/admin/usuarios')
@login_required
@role_required('admin')
def listar_usuarios():
    rows = query("""
        SELECT u.*, r.role_code, r.role_name
        FROM sctis.user_profiles u
        JOIN sctis.user_roles r ON r.role_id = u.role_id
        ORDER BY u.full_name
    """)
    for r in rows:
        r['locked'] = r['locked_until'] is not None and r['locked_until'] > datetime.utcnow()
    return jsonify(rows)


@bp.route('/api/admin/usuarios', methods=['POST'])
@login_required
@role_required('admin')
def crear_usuario():
    data = request.get_json()
    if not data or not data.get('username') or not data.get('password'):
        return jsonify({'error': 'Usuario y contraseña requeridos'}), 400

    username = data['username'].strip().lower()
    existe = query_one("SELECT user_id FROM sctis.user_profiles WHERE username = %s", (username,))
    if existe:
        return jsonify({'error': 'El nombre de usuario ya existe'}), 409

    password_hash = generate_password_hash(data['password'])
    result = query("""
        INSERT INTO sctis.user_profiles
            (username, password_hash, full_name, role_id, estado_codigo, is_active, created_by)
        VALUES (%s, %s, %s, %s, %s, true, %s)
        RETURNING user_id
    """, (
        username,
        password_hash,
        data.get('full_name', '') or username,
        data.get('role_id'),
        data.get('estado_codigo') or None,
        g.user.get('username'),
    ))
    return jsonify({'user_id': result[0]['user_id']}), 201


@bp.route('/api/admin/usuarios/<int:user_id>', methods=['PUT'])
@login_required
@role_required('admin')
def actualizar_usuario(user_id):
    data = request.get_json()
    if not data:
        return jsonify({'error': 'Datos requeridos'}), 400

    user = query_one("SELECT user_id FROM sctis.user_profiles WHERE user_id = %s", (user_id,))
    if not user:
        return jsonify({'error': 'Usuario no encontrado'}), 404

    fields = []
    params = {}
    for key in ('full_name', 'role_id', 'estado_codigo', 'is_active'):
        if key in data:
            fields.append(f"{key} = %({key})s")
            params[key] = data[key]

    if not fields:
        return jsonify({'error': 'Sin campos para actualizar'}), 400

    params['user_id'] = user_id
    query(f"""
        UPDATE sctis.user_profiles
        SET {', '.join(fields)}
        WHERE user_id = %(user_id)s
    """, params, fetch=False)
    return jsonify({'ok': True})


@bp.route('/api/admin/usuarios/<int:user_id>/reset-password', methods=['POST'])
@login_required
@role_required('admin')
def reset_password(user_id):
    data = request.get_json()
    if not data or not data.get('password'):
        return jsonify({'error': 'Nueva contraseña requerida'}), 400

    user = query_one("SELECT user_id FROM sctis.user_profiles WHERE user_id = %s", (user_id,))
    if not user:
        return jsonify({'error': 'Usuario no encontrado'}), 404

    password_hash = generate_password_hash(data['password'])
    query("""
        UPDATE sctis.user_profiles
        SET password_hash = %s, failed_attempts = 0, locked_until = NULL
        WHERE user_id = %s
    """, (password_hash, user_id), fetch=False)
    return jsonify({'ok': True})


# ─── Auditoría de Cargas ──────────────────────────────────────

@bp.route('/admin/auditoria')
@login_required
@role_required('admin')
def auditoria_page():
    return render_template('admin/auditoria.html')


@bp.route('/api/admin/auditoria')
@login_required
@role_required('admin')
def listar_auditoria():
    estado = request.args.get('estado')
    usuario = request.args.get('usuario')
    where = []
    params = []

    if estado:
        where.append("state_code = %s")
        params.append(estado)
    if usuario:
        where.append("ingested_by = %s")
        params.append(usuario)

    where_clause = "WHERE " + " AND ".join(where) if where else ""

    rows = query(f"""
        SELECT submission_id, state_code, ingested_by, ingested_at,
               source_filename, source_sheet, sheet_names,
               row_count, accepted_count, rejected_count,
               validation_status, formato_codigo, correction_file
        FROM audit.submissions
        {where_clause}
        ORDER BY ingested_at DESC
        LIMIT 200
    """, params)
    return jsonify(rows)


@bp.route('/api/admin/auditoria/resumen')
@login_required
@role_required('admin')
def resumen_auditoria():
    rows = query("""
        SELECT
            state_code,
            ingested_by,
            COUNT(*) AS total_cargas,
            SUM(row_count) AS total_registros,
            SUM(accepted_count) AS total_aceptados,
            SUM(rejected_count) AS total_rechazados,
            MAX(ingested_at) AS ultima_carga,
            COUNT(*) FILTER (WHERE validation_status = 'PARTIAL') AS cargas_parciales,
            COUNT(*) FILTER (WHERE validation_status = 'VALIDATED') AS cargas_completas
        FROM audit.submissions
        GROUP BY state_code, ingested_by
        ORDER BY ultima_carga DESC
    """)
    return jsonify(rows)


# ─── Tareas Pendientes ────────────────────────────────────────

@bp.route('/admin/tareas')
@login_required
@role_required('admin')
def tareas_page():
    return render_template('admin/tareas.html')


@bp.route('/api/admin/tareas')
@login_required
@role_required('admin')
def listar_tareas():
    estado = request.args.get('estado')
    estado_tarea = request.args.get('estado_tarea')
    where = ["1=1"]
    params = []

    if estado:
        where.append("tp.estado_codigo = %s")
        params.append(estado)
    if estado_tarea:
        where.append("tp.estado_tarea = %s")
        params.append(estado_tarea)
    else:
        where.append("tp.estado_tarea IN ('PENDIENTE', 'EN_PROCESO')")

    where_clause = " AND ".join(where)

    rows = query(f"""
        SELECT tp.tarea_id, tp.usuario_id, up.username, up.full_name,
               tp.estado_codigo, st.state_name, tp.tipo_tarea,
               tp.descripcion, tp.registros_total, tp.registros_ok,
               tp.registros_rechazados, tp.archivo_correccion,
               tp.estado_tarea, tp.created_at, tp.completed_at,
               tp.horas_limite, tp.alerta_enviada
        FROM sctis.tarea_pendiente tp
        LEFT JOIN sctis.user_profiles up ON up.user_id = tp.usuario_id
        LEFT JOIN common.states st ON st.state_code = tp.estado_codigo
        WHERE {where_clause}
        ORDER BY tp.created_at DESC
    """, params) or []

    now_dt = datetime.now()
    for r in rows:
        created_str = r.get('created_at')
        horas_trans = 0
        vencida = False
        if created_str:
            try:
                if isinstance(created_str, str):
                    c_dt = datetime.strptime(created_str[:19].replace('T', ' '), '%Y-%m-%d %H:%M:%S')
                else:
                    c_dt = created_str
                diff_hours = (now_dt - c_dt).total_seconds() / 3600.0
                horas_trans = round(diff_hours, 1)
                limite = r.get('horas_limite') or 72
                if r.get('estado_tarea') == 'PENDIENTE' and diff_hours > limite:
                    vencida = True
            except Exception:
                pass
        r['horas_transcurridas'] = horas_trans
        r['esta_vencida'] = vencida

    return jsonify(rows)


@bp.route('/api/admin/tareas/<int:tarea_id>/completar', methods=['POST'])
@login_required
@role_required('admin')
def completar_tarea(tarea_id):
    query("""
        UPDATE sctis.tarea_pendiente
        SET estado_tarea = 'COMPLETADA', completed_at = CURRENT_TIMESTAMP
        WHERE tarea_id = %s
    """, (tarea_id,), fetch=False)
    return jsonify({'ok': True})


@bp.route('/api/admin/tareas/<int:tarea_id>/cancelar', methods=['POST'])
@login_required
@role_required('admin')
def cancelar_tarea(tarea_id):
    query("""
        UPDATE sctis.tarea_pendiente
        SET estado_tarea = 'CANCELADA', completed_at = CURRENT_TIMESTAMP
        WHERE tarea_id = %s
    """, (tarea_id,), fetch=False)
    return jsonify({'ok': True})


@bp.route('/api/admin/tareas/correccion/<filename>')
@login_required
@role_required('admin')
def descargar_correccion(filename):
    if '..' in filename or '/' in filename or '\\' in filename:
        return jsonify({'error': 'Nombre de archivo inválido'}), 400
    return send_from_directory(UPLOAD_DIR, filename, as_attachment=True)

# ─── Configuración ────────────────────────────────────────

@bp.route('/admin/configuracion')
@login_required
@role_required('admin')
def configuracion_page():
    return render_template('admin/configuracion.html')

@bp.route('/api/admin/configuracion')
@login_required
@role_required('admin')
def get_configuracion():
    rows = query("SELECT clave, valor, descripcion FROM sctis.configuracion ORDER BY clave")
    return jsonify(rows)

@bp.route('/api/admin/configuracion', methods=['PUT'])
@login_required
@role_required('admin')
def update_configuracion():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'Datos requeridos'}), 400
    
    for clave, valor in data.items():
        query("""
            UPDATE sctis.configuracion
            SET valor = %s, updated_at = now(), updated_by = %s
            WHERE clave = %s
        """, (str(valor), g.user.get('username'), clave), fetch=False)
        
    return jsonify({'ok': True})

from app.workspace_integration import upload_to_drive, send_notification_email
import pandas as pd
from docx import Document
from datetime import datetime

@bp.route('/api/admin/generar_reporte', methods=['POST'])
@login_required
@role_required('admin')
def generar_reporte_ws():
    data = request.get_json()
    access_token = data.get('access_token')
    if not access_token:
        return jsonify({'error': 'No OAuth token provided'}), 400
        
    try:
        # Generar datos
        resumen = query("""
            SELECT st.state_name as estado, 
                   COUNT(t.tira_id) as total_interrupciones,
                   COALESCE(SUM(t.horas), 0) as total_horas,
                   COALESCE(SUM(t.kva), 0) as total_kva
            FROM sctis.tira_interrupcion t
            JOIN common.states st ON st.state_code = t.estado_codigo
            GROUP BY st.state_name
        """)
        
        # 1. Crear documento docx
        doc_path = f"/tmp/resumen_gestion_{datetime.now().strftime('%Y%m%d')}.docx"
        doc = Document()
        doc.add_heading('Resumen de Gestión de Interrupciones', 0)
        doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Estado'
        hdr_cells[1].text = 'Interrupciones'
        hdr_cells[2].text = 'Horas Totales'
        hdr_cells[3].text = 'KVA Afectados'
        
        for r in resumen:
            row_cells = table.add_row().cells
            row_cells[0].text = r['estado']
            row_cells[1].text = str(r['total_interrupciones'])
            row_cells[2].text = f"{r['total_horas']:.2f}"
            row_cells[3].text = f"{r['total_kva']:.2f}"
            
        doc.save(doc_path)
        
        # 2. Crear excel
        excel_path = f"/tmp/datos_resumen_{datetime.now().strftime('%Y%m%d')}.xlsx"
        df = pd.DataFrame(resumen)
        df.to_excel(excel_path, index=False)
        
        # 3. Subir a Google Drive
        folder_id = '1FCNvnPrCPzV4VKNfcUcN2sZC88XS3kzG'
        doc_id = upload_to_drive(access_token, doc_path, folder_id, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', f"Resumen_{datetime.now().strftime('%Y%m%d')}.docx")
        xls_id = upload_to_drive(access_token, excel_path, folder_id, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', f"Datos_{datetime.now().strftime('%Y%m%d')}.xlsx")
        
        # 4. Enviar email de notificación al admin (opcional)
        email_config = query_one("SELECT valor FROM sctis.configuracion WHERE clave = 'notificacion_email'")
        if email_config and email_config['valor']:
            send_notification_email(access_token, email_config['valor'], "Reporte de Gestión Generado", f"Se han generado y subido los reportes a Google Drive.\nDoc ID: {doc_id}\nExcel ID: {xls_id}")
            
        return jsonify({
            'ok': True,
            'doc_id': doc_id,
            'xls_id': xls_id
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── MÓDULO DE DESCARGA DE DATOS DE INTERRUPCIONES (SOLO ADMINS) ──────────

COLUMNAS_EXCEL_EXPORT = [
    ('ESTADO', 15),
    ('SISTEMA', 12),
    ('DISTRITO', 12),
    ('SUBESTACIÓN', 20),
    ('CIRCUITO', 15),
    ('FECHA INICIO', 14),
    ('HORA INICIO', 10),
    ('FECHA FIN', 14),
    ('HORA FIN', 10),
    ('DURACIÓN', 10),
    ('CARGA', 8),
    ('FREC', 6),
    ('HORAS', 8),
    ('TTI CTO', 8),
    ('SEÑAL', 6),
    ('CAUSA', 20),
    ('SUB-CAUSA', 20),
    ('OBSERVACIÓN', 30),
    ('SECTORES', 25),
    ('CIUDAD', 15),
    ('KVA', 8),
]


def _build_download_query(filters):
    where = ["ti.activo = true"]
    params = []

    anio = filters.get('anio')
    if anio and str(anio).strip() and str(anio).upper() != 'TODOS':
        where.append("""(
            EXTRACT(YEAR FROM ti.fecha_falla)::text = %s 
            OR strftime('%%Y', ti.fecha_falla) = %s 
            OR ti.fecha_falla::text LIKE %s
            OR ti.fecha_inicio::text LIKE %s
        )""")
        params.extend([str(anio), str(anio), f"{anio}-%", f"{anio}-%"])

    mes = filters.get('mes')
    if mes and str(mes).strip() and str(mes).upper() != 'TODOS':
        where.append("ti.mes = %s")
        params.append(str(mes).upper())

    estado = filters.get('estado')
    if estado and str(estado).strip() and str(estado).upper() != 'TODOS':
        where.append("ti.estado_codigo = %s")
        params.append(str(estado))

    causa_id = filters.get('causa_id')
    if causa_id and str(causa_id).strip() and str(causa_id).upper() != 'TODAS':
        where.append("ti.causa_id = %s")
        params.append(int(causa_id))

    subestacion = filters.get('subestacion')
    if subestacion and str(subestacion).strip() and str(subestacion).upper() != 'TODAS':
        if str(subestacion).isdigit():
            where.append("ti.subestacion_id = %s")
            params.append(int(subestacion))
        else:
            where.append("LOWER(ti.subestacion) = LOWER(%s)")
            params.append(str(subestacion))

    circuito = filters.get('circuito')
    if circuito and str(circuito).strip() and str(circuito).upper() != 'TODOS':
        if str(circuito).isdigit():
            where.append("ti.circuito_id = %s")
            params.append(int(circuito))
        else:
            where.append("LOWER(ti.circuito) = LOWER(%s)")
            params.append(str(circuito))

    where_sql = " WHERE " + " AND ".join(where)

    sql = f"""
        SELECT 
            ti.tira_id, ti.estado_codigo, s.state_name, ti.sistema, ti.jefatura,
            ti.subestacion, ti.circuito, ti.fecha_inicio, ti.fecha_fin,
            ti.duracion_calculada, ti.horas, ti.mw, ti.causa, ti.sub_causa,
            ti.observacion, ti.sectores, ti.ciudad, ti.kva, ti.es_excepcion_admin,
            c.causa_nombre AS causa_homologada,
            COALESCE(se.asset_name_normalizado, se.asset_name) AS subestacion_normalizado,
            COALESCE(ci.asset_name_normalizado, ci.asset_name) AS circuito_normalizado
        FROM sctis.tira_interrupcion ti
        LEFT JOIN common.states s ON s.state_code = ti.estado_codigo
        LEFT JOIN sctis.causa c ON c.causa_id = ti.causa_id
        LEFT JOIN common.assets se ON se.asset_id = ti.subestacion_id
        LEFT JOIN common.assets ci ON ci.asset_id = ti.circuito_id
        {where_sql}
        ORDER BY ti.fecha_falla DESC, ti.tira_id DESC
    """
    return sql, params


def _format_record_row(r):
    estado_val = r.get('state_name') or r.get('estado_codigo') or ''
    sistema_val = r.get('sistema') or 'DISTRIBUCION'
    distrito_val = r.get('jefatura') or ''
    se_val = r.get('subestacion_normalizado') or r.get('subestacion') or ''
    ci_val = r.get('circuito_normalizado') or r.get('circuito') or ''

    f_ini_val = r.get('fecha_inicio')
    f_fin_val = r.get('fecha_fin')

    f_ini_str, h_ini_str = '', ''
    if f_ini_val:
        if isinstance(f_ini_val, str):
            parts = f_ini_val.split('T') if 'T' in f_ini_val else f_ini_val.split(' ')
            f_ini_str = parts[0] if parts else ''
            h_ini_str = parts[1][:5] if len(parts) > 1 else ''
        elif isinstance(f_ini_val, (datetime, date)):
            f_ini_str = f_ini_val.strftime('%Y-%m-%d')
            h_ini_str = f_ini_val.strftime('%H:%M') if hasattr(f_ini_val, 'strftime') else ''

    f_fin_str, h_fin_str = '', ''
    if f_fin_val:
        if isinstance(f_fin_val, str):
            parts = f_fin_val.split('T') if 'T' in f_fin_val else f_fin_val.split(' ')
            f_fin_str = parts[0] if parts else ''
            h_fin_str = parts[1][:5] if len(parts) > 1 else ''
        elif isinstance(f_fin_val, (datetime, date)):
            f_fin_str = f_fin_val.strftime('%Y-%m-%d')
            h_fin_str = f_fin_val.strftime('%H:%M') if hasattr(f_fin_val, 'strftime') else ''

    duracion_val = r.get('duracion_calculada') or ''
    carga_val = r.get('mw') if r.get('mw') is not None else 0
    frec_val = 1
    horas_val = r.get('horas') if r.get('horas') is not None else 0
    tti_cto_val = 0
    senal_val = 0
    causa_val = r.get('causa_homologada') or r.get('causa') or ''
    sub_causa_val = r.get('sub_causa') or ''
    obs_val = r.get('observacion') or ''
    sectores_val = r.get('sectores') or ''
    ciudad_val = r.get('ciudad') or ''
    kva_val = r.get('kva') if r.get('kva') is not None else 0

    return [
        estado_val, sistema_val, distrito_val, se_val, ci_val,
        f_ini_str, h_ini_str, f_fin_str, h_fin_str, duracion_val,
        carga_val, frec_val, horas_val, tti_cto_val, senal_val,
        causa_val, sub_causa_val, obs_val, sectores_val, ciudad_val, kva_val
    ]


@bp.route('/admin/descargas')
@login_required
@role_required('admin')
def descargas_page():
    return render_template('admin/descargas.html')


@bp.route('/api/admin/descargas/filtros-opciones')
@login_required
@role_required('admin')
def descargas_filtros_opciones():
    fechas_rows = query("SELECT DISTINCT fecha_falla FROM sctis.tira_interrupcion WHERE fecha_falla IS NOT NULL") or []
    anios_set = set()
    for r in fechas_rows:
        f = str(r.get('fecha_falla') or '').strip()
        if len(f) >= 4 and f[:4].isdigit():
            anios_set.add(int(f[:4]))
    anios = sorted(list(anios_set), reverse=True)
    if not anios:
        anios = [datetime.now().year]

    estados = query("SELECT state_code, state_name FROM common.states ORDER BY state_name")
    causas = query("SELECT causa_id, causa_codigo, causa_nombre FROM sctis.causa ORDER BY causa_nombre")

    subestaciones = query("""
        SELECT asset_id, COALESCE(asset_name_normalizado, asset_name) AS asset_name, state_code
        FROM common.assets
        WHERE asset_type = 'SUBSTATION' AND is_active = true
        ORDER BY asset_name
    """)

    circuitos = query("""
        SELECT asset_id, COALESCE(asset_name_normalizado, asset_name) AS asset_name, parent_asset_id, state_code
        FROM common.assets
        WHERE asset_type = 'CIRCUITO' AND is_active = true
        ORDER BY asset_name
    """)

    meses = ['ENERO', 'FEBRERO', 'MARZO', 'ABRIL', 'MAYO', 'JUNIO', 'JULIO', 'AGOSTO', 'SEPTIEMBRE', 'OCTUBRE', 'NOVIEMBRE', 'DICIEMBRE']

    return jsonify({
        'anios': anios,
        'meses': meses,
        'estados': estados or [],
        'causas': causas or [],
        'subestaciones': subestaciones or [],
        'circuitos': circuitos or []
    })


@bp.route('/api/admin/descargas/preview', methods=['POST'])
@login_required
@role_required('admin')
def descargas_preview():
    filters = request.get_json() or {}
    sql, params = _build_download_query(filters)
    rows = query(sql, params) or []

    total = len(rows)
    total_horas = sum(r.get('horas') or 0 for r in rows)
    total_kva = sum(r.get('kva') or 0 for r in rows)
    excepciones_count = sum(1 for r in rows if r.get('es_excepcion_admin'))

    sample = [_format_record_row(r) for r in rows[:10]]

    return jsonify({
        'total_registros': total,
        'total_horas': round(total_horas, 2),
        'total_kva': round(total_kva, 2),
        'excepciones_count': excepciones_count,
        'sample_rows': sample
    })


@bp.route('/api/admin/descargas/exportar', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def descargas_exportar():
    if request.method == 'POST':
        filters = request.get_json() or {}
    else:
        filters = {
            'anio': request.args.get('anio'),
            'mes': request.args.get('mes'),
            'estado': request.args.get('estado'),
            'causa_id': request.args.get('causa_id'),
            'subestacion': request.args.get('subestacion'),
            'circuito': request.args.get('circuito'),
            'formato': request.args.get('formato', 'EXCEL').upper()
        }

    formato = (filters.get('formato') or 'EXCEL').upper()
    sql, params = _build_download_query(filters)
    rows = query(sql, params) or []

    user_id = session.get('user_id') or session.get('user', {}).get('user_id')
    username = session.get('user', {}).get('username') or session.get('username') or 'admin'
    ip_addr = request.remote_addr or '127.0.0.1'
    user_agent = request.headers.get('User-Agent', '')

    audit_sql = """
        INSERT INTO "sctis.audit_descarga_datos"
            (user_id, username, formato, filtros_json, total_registros, ip_address, user_agent)
        VALUES (%s, %s, %s, %s, %s, %s, %s)
    """
    try:
        query(audit_sql, (
            user_id, username, formato,
            json.dumps(filters, ensure_ascii=False),
            len(rows), ip_addr, user_agent
        ), fetch=False)
    except Exception as e:
        print("Error logging download audit:", e)

    timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')

    if formato == 'CSV':
        buf = io.StringIO()
        buf.write('\ufeff')
        writer = csv.writer(buf, delimiter=';', quoting=csv.QUOTE_MINIMAL)

        headers = [col[0] for col in COLUMNAS_EXCEL_EXPORT]
        writer.writerow(headers)

        for r in rows:
            writer.writerow(_format_record_row(r))

        output_bytes = io.BytesIO(buf.getvalue().encode('utf-8-sig'))
        output_bytes.seek(0)
        filename = f"SCTIS_Interrupciones_{timestamp_str}.csv"

        return send_file(
            output_bytes,
            mimetype='text/csv; charset=utf-8',
            as_attachment=True,
            download_name=filename
        )

    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'FORMULARIO DE CARGA'

        hdr_font = Font(name='Calibri', bold=True, color='FFFFFF', size=10)
        hdr_fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
        hdr_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )

        for col_idx, (header, width) in enumerate(COLUMNAS_EXCEL_EXPORT, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = hdr_align
            cell.border = thin_border
            ws.column_dimensions[get_column_letter(col_idx)].width = width

        for row_idx, r in enumerate(rows, 2):
            row_data = _format_record_row(r)
            for col_idx, val in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=val)
                cell.border = thin_border

        ws.freeze_panes = 'A2'

        ws2 = wb.create_sheet('INSTRUCTIVO DE CARGA')
        title_font = Font(name='Calibri', bold=True, size=14, color='1F4E79')
        subtitle_font = Font(name='Calibri', bold=True, size=11, color='1F4E79')
        body_font = Font(name='Calibri', size=10)

        ws2.cell(row=1, column=1, value='EXPORTACIÓN DE DATOS - FORMULARIO OFICIAL SCTIS').font = title_font
        ws2.column_dimensions['A'].width = 5
        ws2.column_dimensions['B'].width = 90

        instrucciones = [
            ('1. DESCRIPCIÓN DEL ARCHIVO', [
                f'Archivo generado por el Sistema SCTIS el {datetime.now().strftime("%d/%m/%Y a las %H:%M:%S")}.',
                f'Total de registros exportados: {len(rows)}.',
                f'Usuario autorizante: {username}.',
            ]),
            ('2. CUMPLIMIENTO NORMAS ISO 27001 / ISO 8000', [
                'Este libro contiene la estructura oficial homologada de 21 columnas de datos.',
                'Los datos exportados pueden ser re-importados directamente en el sistema.',
            ]),
            ('3. AUDITORÍA Y NO REPUDIO', [
                'La descarga de esta información ha sido registrada en el Libro de Auditoría de Descargas del Sistema.',
            ])
        ]

        r_idx = 3
        for tit, items in instrucciones:
            ws2.cell(row=r_idx, column=2, value=tit).font = subtitle_font
            r_idx += 1
            for item in items:
                ws2.cell(row=r_idx, column=2, value=item).font = body_font
                r_idx += 1
            r_idx += 1

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        filename = f"SCTIS_Interrupciones_{timestamp_str}.xlsx"

        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )


@bp.route('/api/admin/descargas/auditoria')
@login_required
@role_required('admin')
def descargas_auditoria():
    rows = query("""
        SELECT download_id, user_id, username, formato, filtros_json,
               total_registros, ip_address, created_at
        FROM "sctis.audit_descarga_datos"
        ORDER BY created_at DESC
        LIMIT 100
    """)
    return jsonify(rows or [])


# ─── MÓDULO DE MONITOREO DE OPORTUNIDAD Y CONTROL DE CARGAS (ING. FAVIO) ───

@bp.route('/admin/monitoreo-cargas')
@login_required
@role_required('admin')
def monitoreo_cargas_page():
    return render_template('admin/monitoreo.html')


@bp.route('/api/admin/monitoreo-cargas/matriz')
@login_required
@role_required('admin')
def monitoreo_cargas_matriz():
    estados = query("SELECT state_code, state_name FROM common.states ORDER BY state_name") or []
    
    # 1. Audit Submissions por estado
    last_sub_rows = query("""
        SELECT state_code, MAX(ingested_at) AS ultima_carga, COUNT(*) AS total_cargas, SUM(row_count) AS total_filas
        FROM audit.submissions
        GROUP BY state_code
    """) or []
    subs_map = {r['state_code']: r for r in last_sub_rows}

    # 2. Audit Admin cargas excepcionales
    last_admin_rows = query("""
        SELECT estado_codigo, MAX(created_at) AS ultima_carga, COUNT(*) AS total_cargas, SUM(total_registros) AS total_filas
        FROM sctis.audit_admin_carga_excepcional
        GROUP BY estado_codigo
    """) or []
    admin_map = {r['estado_codigo']: r for r in last_admin_rows}

    # 3. Tira interrupcion registros directos
    tira_rows = query("""
        SELECT estado_codigo, MAX(created_at) AS ultima_carga, COUNT(*) AS total_filas
        FROM sctis.tira_interrupcion
        WHERE estado_codigo IS NOT NULL
        GROUP BY estado_codigo
    """) or []
    tira_map = {r['estado_codigo']: r for r in tira_rows}

    # Incompletos por estado
    incompletos_rows = query("""
        SELECT estado_codigo, COUNT(*) AS total_incompletos
        FROM sctis.tira_interrupcion
        WHERE es_excepcion_admin = 1 OR estado_calculo = 'INCOMPLETO_EXCEPCION_ADMIN'
        GROUP BY estado_codigo
    """) or []
    inc_map = {r['estado_codigo']: r['total_incompletos'] for r in incompletos_rows}

    # Tareas por estado
    tareas_rows = query("""
        SELECT estado_codigo,
               COUNT(*) AS total_tareas,
               COUNT(CASE WHEN estado_tarea IN ('PENDIENTE', 'EN_PROCESO') THEN 1 END) AS tareas_activas
        FROM sctis.tarea_pendiente
        GROUP BY estado_codigo
    """) or []
    tareas_map = {r['estado_codigo']: r for r in tareas_rows}

    now = datetime.now()
    dia_semana = now.weekday()  # 0=Lunes, 2=Miércoles, 3=Jueves, 6=Domingo
    dia_mes = now.day

    matriz = []

    for st in estados:
        sc = st['state_code']
        sname = st['state_name']
        
        # Buscar la última carga de todas las fuentes disponibles
        candidates = []
        c_sub = subs_map.get(sc, {})
        if c_sub.get('ultima_carga'):
            candidates.append(str(c_sub['ultima_carga']))
        c_adm = admin_map.get(sc, {})
        if c_adm.get('ultima_carga'):
            candidates.append(str(c_adm['ultima_carga']))
        c_tir = tira_map.get(sc, {})
        if c_tir.get('ultima_carga'):
            candidates.append(str(c_tir['ultima_carga']))

        ult_carga = max(candidates) if candidates else None
        
        incompletos = inc_map.get(sc, 0)
        t_info = tareas_map.get(sc, {})
        tareas_activas = t_info.get('tareas_activas', 0)

        # Evaluación Semanal (Miércoles / Jueves - Carga recibida hace <= 4 días)
        dias_desde_carga = 999
        if ult_carga:
            try:
                d_obj = datetime.strptime(ult_carga[:19].replace('T', ' '), '%Y-%m-%d %H:%M:%S')
                dias_desde_carga = (now - d_obj).days
            except Exception:
                dias_desde_carga = 999

        cumplio_semana = (dias_desde_carga <= 4)
        if cumplio_semana:
            estatus_semanal = 'AL_DIA'
            estatus_semanal_txt = 'Al Día'
        elif dia_semana in (2, 3) and dias_desde_carga > 4:
            estatus_semanal = 'EN_RIESGO'
            estatus_semanal_txt = 'Carga Requerida Hoy (Mié/Jue)'
        else:
            estatus_semanal = 'MOROSO'
            estatus_semanal_txt = 'Falta Carga Semanal'

        # Evaluación Mensual (Consolidado - Carga recibida hace <= 25 días)
        cumplio_mes = (dias_desde_carga <= 25)
        if cumplio_mes:
            estatus_mensual = 'AL_DIA'
            estatus_mensual_txt = 'Consolidado Al Día'
        elif dia_mes <= 3:
            estatus_mensual = 'EN_RIESGO'
            estatus_mensual_txt = 'En Plazo Mensual (Hasta Día 3)'
        else:
            estatus_mensual = 'MOROSO'
            estatus_mensual_txt = 'Falta Consolidado Mensual'

        tot_filas = (c_sub.get('total_filas') or 0) + (c_adm.get('total_filas') or 0) + (c_tir.get('total_filas') or 0)

        matriz.append({
            'state_code': sc,
            'state_name': sname,
            'ultima_carga': ult_carga or 'SIN CARGAS REGISTRADAS',
            'total_filas': tot_filas,
            'incompletos_count': incompletos,
            'tareas_activas': tareas_activas,
            'cumplio_semana': cumplio_semana,
            'estatus_semanal': estatus_semanal,
            'estatus_semanal_txt': estatus_semanal_txt,
            'cumplio_mes': cumplio_mes,
            'estatus_mensual': estatus_mensual,
            'estatus_mensual_txt': estatus_mensual_txt,
            'dias_desde_carga': dias_desde_carga if dias_desde_carga < 999 else None
        })

    # Resumen y Listas de Control Rápido
    cumplieron_semana_list = [m for m in matriz if m['cumplio_semana']]
    faltan_semana_list = [m for m in matriz if not m['cumplio_semana']]

    cumplieron_mes_list = [m for m in matriz if m['cumplio_mes']]
    faltan_mes_list = [m for m in matriz if not m['cumplio_mes']]

    total_al_dia = sum(1 for m in matriz if m['cumplio_semana'] and m['cumplio_mes'])
    total_morosos = sum(1 for m in matriz if m['estatus_semanal'] == 'MOROSO' or m['estatus_mensual'] == 'MOROSO')
    total_riesgo = len(matriz) - total_al_dia - total_morosos
    total_incompletos_global = sum(m['incompletos_count'] for m in matriz)

    return jsonify({
        'fecha_evaluacion': now.strftime('%Y-%m-%d %H:%M:%S'),
        'dia_semana_num': dia_semana,
        'dia_mes_num': dia_mes,
        'resumen': {
            'total_estados': len(matriz),
            'total_al_dia': total_al_dia,
            'total_riesgo': total_riesgo,
            'total_morosos': total_morosos,
            'total_incompletos_global': total_incompletos_global,
            'cant_cumplieron_semana': len(cumplieron_semana_list),
            'cant_faltan_semana': len(faltan_semana_list),
            'cant_cumplieron_mes': len(cumplieron_mes_list),
            'cant_faltan_mes': len(faltan_mes_list)
        },
        'listas': {
            'faltan_semana': [{'code': m['state_code'], 'name': m['state_name'], 'dias': m['dias_desde_carga']} for m in faltan_semana_list],
            'faltan_mes': [{'code': m['state_code'], 'name': m['state_name'], 'dias': m['dias_desde_carga']} for m in faltan_mes_list],
            'cumplieron_semana': [{'code': m['state_code'], 'name': m['state_name']} for m in cumplieron_semana_list],
            'cumplieron_mes': [{'code': m['state_code'], 'name': m['state_name']} for m in cumplieron_mes_list]
        },
        'matriz': matriz
    })


@bp.route('/api/admin/monitoreo-cargas/crear-tarea', methods=['POST'])
@login_required
@role_required('admin')
def crear_tarea_saneamiento():
    data = request.get_json() or {}
    estado_codigo = data.get('estado_codigo')
    tipo_tarea = data.get('tipo_tarea', 'REQUERIMIENTO_DATA_INCOMPLETA')
    descripcion = data.get('descripcion')
    horas_limite = int(data.get('horas_limite', 48))

    if not estado_codigo or not descripcion:
        return jsonify({'error': 'Estado y descripción son obligatorios'}), 400

    # Buscar usuario operador principal del estado o asignarlo a admin
    user = query_one("""
        SELECT user_id FROM sctis.user_profiles
        WHERE estado_codigo = %s AND is_active = true
        LIMIT 1
    """, (estado_codigo,))
    
    usuario_id = user['user_id'] if user else (session.get('user_id') or 1)

    # Incompletos actuales del estado
    inc_row = query_one("""
        SELECT COUNT(*) AS total
        FROM sctis.tira_interrupcion
        WHERE estado_codigo = %s AND (es_excepcion_admin = 1 OR estado_calculo = 'INCOMPLETO_EXCEPCION_ADMIN')
    """, (estado_codigo,))
    
    total_inc = inc_row['total'] if inc_row else 0

    sql = """
        INSERT INTO sctis.tarea_pendiente
            (usuario_id, estado_codigo, tipo_tarea, descripcion, registros_total,
             estado_tarea, horas_limite)
        VALUES (%s, %s, %s, %s, %s, 'PENDIENTE', %s)
    """
    query(sql, (usuario_id, estado_codigo, tipo_tarea, descripcion, total_inc, horas_limite), fetch=False)

    return jsonify({'ok': True, 'message': 'Tarea asignada exitosamente al estado'})

