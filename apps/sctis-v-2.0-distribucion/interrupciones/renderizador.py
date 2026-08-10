"""Renderer que convierte bloques de contenido en Markdown (.md) y DOCX (.docx).

Permite mantener una única fuente de contenido por documento y generar ambos
formatos con la misma información, aprovechando los helpers de docx_utils.py.
"""
from docx.shared import Pt
from docx_utils import (
    crear_documento, add_portada, add_control_documento, add_flow_table,
    add_branching_flow, add_info_box, add_tech_note, add_decision_table,
    add_simple_table, agregar_glosario, pie_de_documento, guardar,
)


def _md_table(headers, rows):
    sep = "| " + " | ".join(["---"] * len(headers)) + " |"
    header = "| " + " | ".join(headers) + " |"
    lines = [header, sep]
    for r in rows:
        lines.append("| " + " | ".join(str(c) for c in r) + " |")
    return "\n".join(lines)


def render_markdown(blocks, titulo, subtitulo, sistema):
    """Renderiza los bloques a texto Markdown completo."""
    out = []
    out.append(f"# {titulo}\n")
    out.append(f"**{subtitulo}**  \n**{sistema}**  \n")
    out.append("---\n")
    for block in blocks:
        kind = block[0]
        if kind == "h1":
            out.append(f"\n## {block[1]}\n")
        elif kind == "h2":
            out.append(f"\n### {block[1]}\n")
        elif kind == "h3":
            out.append(f"\n#### {block[1]}\n")
        elif kind == "p":
            for p in block[1].split("\n\n"):
                out.append(p.strip() + "\n")
        elif kind == "bullets":
            for item in block[1]:
                out.append(f"- {item}")
            out.append("")
        elif kind == "flow":
            pasos = " → ".join(t for t, _ in block[1])
            out.append(f"> Flujo: **{pasos}**")
            for t, d in block[1]:
                out.append(f"> - **{t}**: {d}")
            out.append("")
        elif kind == "branch":
            out.append(f"**{block[1]}**  ")
            out.append(f"*{block[2]} ({block[4]}):* " + "; ".join(block[2]))
            out.append(f"*{block[3]} ({block[5]}):* " + "; ".join(block[3]))
            out.append("")
        elif kind == "decision":
            out.append(_md_table(block[2], block[1]))
            out.append("")
        elif kind == "table":
            out.append(_md_table(block[1], block[2]))
            out.append("")
        elif kind == "info":
            out.append(f"> **{block[2]}:** {block[1]}\n")
        elif kind == "tech":
            out.append(f"> **{block[2]}:** {block[1]}\n")
        elif kind == "glossary":
            out.append(_md_table(["Término", "Definición"], block[1]))
            out.append("")
        elif kind == "code":
            out.append("```\n" + block[1] + "\n```")
            out.append("")
        elif kind == "pagebreak":
            out.append("\n---\n")
    out.append("\n---\n*Fin del documento.*")
    return "\n".join(out)


def render_doc(blocks, titulo, subtitulo, sistema, doc_codigo, ruta, fecha="Agosto 2026",
               extra_pie=""):
    """Renderiza los bloques a un archivo .doc (HTML estructurado) 100% compatible con Google Drive y Google Docs."""
    parts = []
    parts.append('<!DOCTYPE html>')
    parts.append('<html xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:w="urn:schemas-microsoft-com:office:word" xmlns="http://www.w3.org/TR/REC-html40">')
    parts.append('<head>')
    parts.append('<meta charset="utf-8">')
    parts.append(f'<title>{titulo}</title>')
    parts.append('<style>')
    parts.append("body { font-family: Arial, sans-serif; font-size: 11pt; color: #1e293b; line-height: 1.6; margin: 40px; }")
    parts.append("h1 { color: #003C71; font-size: 20pt; margin-top: 24pt; margin-bottom: 12pt; border-bottom: 2px solid #003C71; padding-bottom: 6px; }")
    parts.append("h2 { color: #003C71; font-size: 15pt; margin-top: 18pt; margin-bottom: 8pt; }")
    parts.append("h3 { color: #003C71; font-size: 12pt; margin-top: 14pt; margin-bottom: 6pt; }")
    parts.append("p { margin-bottom: 10pt; text-align: justify; }")
    parts.append("table { border-collapse: collapse; width: 100%; margin: 14pt 0; }")
    parts.append("th, td { border: 1px solid #cbd5e1; padding: 8pt 10pt; text-align: left; vertical-align: top; font-size: 10pt; }")
    parts.append("th { background-color: #003C71; color: #ffffff; font-weight: bold; }")
    parts.append("tr:nth-child(even) td { background-color: #f8fafc; }")
    parts.append(".portada { text-align: center; padding: 40px 20px; border: 2px solid #003C71; background-color: #f8fafc; margin-bottom: 30px; border-radius: 8px; }")
    parts.append(".portada h1 { border-bottom: none; font-size: 24pt; margin-bottom: 12pt; color: #003C71; }")
    parts.append(".portada .subtitulo { font-size: 13pt; color: #334155; font-weight: bold; margin-bottom: 8pt; }")
    parts.append(".portada .sistema { font-size: 11pt; color: #64748b; margin-bottom: 16pt; }")
    parts.append(".control-doc { width: 100%; margin-bottom: 24pt; border: 1px solid #cbd5e1; }")
    parts.append(".control-doc td { padding: 6pt 10pt; font-size: 9.5pt; }")
    parts.append(".control-doc .lbl { font-weight: bold; background-color: #e2e8f0; width: 25%; color: #1e293b; }")
    parts.append(".info-box { background-color: #f0f9ff; border-left: 4px solid #0284c7; padding: 10pt 14pt; margin: 12pt 0; border-radius: 4px; }")
    parts.append(".tech-box { background-color: #f8fafc; border-left: 4px solid #475569; padding: 10pt 14pt; margin: 12pt 0; border-radius: 4px; }")
    parts.append(".box-title { font-weight: bold; margin-bottom: 4pt; color: #0f172a; }")
    parts.append("ul, ol { margin-bottom: 10pt; padding-left: 24pt; }")
    parts.append("li { margin-bottom: 4pt; }")
    parts.append("pre { background-color: #f1f5f9; padding: 10pt; border-radius: 4px; border: 1px solid #e2e8f0; font-size: 9pt; font-family: monospace; }")
    parts.append(".footer { margin-top: 40pt; border-top: 1px solid #cbd5e1; padding-top: 10pt; text-align: center; font-size: 9pt; color: #64748b; }")
    parts.append('</style>')
    parts.append('</head>')
    parts.append('<body>')

    # Portada
    parts.append(f'<div class="portada"><h1>{titulo}</h1><div class="subtitulo">{subtitulo}</div><div class="sistema">{sistema}</div><div class="fecha">{fecha}</div></div>')

    # Control doc
    parts.append(f'<table class="control-doc"><tr><td class="lbl">Código Documento:</td><td>{doc_codigo}</td><td class="lbl">Fecha:</td><td>{fecha}</td></tr><tr><td class="lbl">Sistema:</td><td>{sistema}</td><td class="lbl">Estado:</td><td>VIGENTE / APROBADO</td></tr></table>')

    for block in blocks:
        kind = block[0]
        if kind == "h1":
            parts.append(f'<h1>{block[1]}</h1>')
        elif kind == "h2":
            parts.append(f'<h2>{block[1]}</h2>')
        elif kind == "h3":
            parts.append(f'<h3>{block[1]}</h3>')
        elif kind == "p":
            for p in block[1].split("\n\n"):
                parts.append(f'<p>{p.strip()}</p>')
        elif kind == "bullets":
            parts.append('<ul>')
            for item in block[1]:
                parts.append(f'<li>{item}</li>')
            parts.append('</ul>')
        elif kind == "flow":
            parts.append('<table><thead><tr><th>Paso</th><th>Descripción</th></tr></thead><tbody>')
            for paso, desc in block[1]:
                parts.append(f'<tr><td><b>{paso}</b></td><td>{desc}</td></tr>')
            parts.append('</tbody></table>')
        elif kind == "branch":
            parts.append(f'<h3>{block[1]}</h3>')
            parts.append('<table><thead><tr><th>Condición</th><th>Acciones</th></tr></thead><tbody>')
            parts.append(f'<tr><td><b>{block[4]}</b></td><td>' + '<br>'.join(block[2]) + '</td></tr>')
            parts.append(f'<tr><td><b>{block[5]}</b></td><td>' + '<br>'.join(block[3]) + '</td></tr>')
            parts.append('</tbody></table>')
        elif kind in ("decision", "table"):
            headers = block[2] if kind == "decision" else block[1]
            rows = block[1] if kind == "decision" else block[2]
            parts.append('<table><thead><tr>')
            for h in headers:
                parts.append(f'<th>{h}</th>')
            parts.append('</tr></thead><tbody>')
            for r in rows:
                parts.append('<tr>')
                for cell in r:
                    parts.append(f'<td>{cell}</td>')
                parts.append('</tr>')
            parts.append('</tbody></table>')
        elif kind in ("info", "tech"):
            css_cls = "info-box" if kind == "info" else "tech-box"
            parts.append(f'<div class="{css_cls}"><div class="box-title">{block[2]}</div><p>{block[1]}</p></div>')
        elif kind == "glossary":
            parts.append('<table><thead><tr><th>Término</th><th>Definición</th></tr></thead><tbody>')
            for term, defn in block[1]:
                parts.append(f'<tr><td><b>{term}</b></td><td>{defn}</td></tr>')
            parts.append('</tbody></table>')
        elif kind == "code":
            parts.append(f'<pre>{block[1]}</pre>')
        elif kind == "pagebreak":
            parts.append('<hr style="margin: 30px 0; border: none; border-top: 1px dashed #cbd5e1;">')

    parts.append(f'<div class="footer"><p>{doc_codigo} - {titulo} | {sistema} {extra_pie}</p></div></body></html>')

    with open(ruta, 'w', encoding='utf-8') as f:
        f.write('\n'.join(parts))


def render_docx(blocks, titulo, subtitulo, sistema, doc_codigo, ruta, fecha="Agosto 2026",
                extra_pie=""):
    """Renderiza los bloques a un DOCX con formato ISO."""
    doc = crear_documento()
    add_portada(doc, titulo, subtitulo, sistema=sistema, fecha=fecha)
    add_control_documento(doc, doc_codigo, titulo, fecha=fecha)

    for block in blocks:
        kind = block[0]
        if kind == "h1":
            doc.add_heading(block[1], level=1)
        elif kind == "h2":
            doc.add_heading(block[1], level=2)
        elif kind == "h3":
            doc.add_heading(block[1], level=3)
        elif kind == "p":
            for p in block[1].split("\n\n"):
                doc.add_paragraph(p.strip())
        elif kind == "bullets":
            for item in block[1]:
                doc.add_paragraph(item, style='List Bullet')
        elif kind == "flow":
            add_flow_table(doc, block[1])
        elif kind == "branch":
            add_branching_flow(doc, block[1], block[2], block[3],
                               left_label=block[4], right_label=block[5])
        elif kind == "decision":
            add_decision_table(doc, block[1], block[2])
        elif kind == "table":
            colw = block[3] if len(block) > 3 else None
            add_simple_table(doc, block[2], block[1], col_widths=colw)
        elif kind == "info":
            add_info_box(doc, block[1], label=block[2])
        elif kind == "tech":
            add_tech_note(doc, block[1], label=block[2])
        elif kind == "glossary":
            agregar_glosario(doc, block[1])
        elif kind == "code":
            for linea in block[1].split("\n"):
                p = doc.add_paragraph()
                run = p.add_run(linea)
                run.font.name = 'Consolas'
                run.font.size = Pt(8)
        elif kind == "pagebreak":
            doc.add_page_break()

    pie_de_documento(doc, extra_pie)
    guardar(doc, ruta)
