"""Utilidades compartidas para generar documentos DOCX estilo ISO.

Reutiliza los helpers de estilo de generar_documento.py (portada, flujos,
cajas informativas, tablas de decisión) para mantener consistencia visual
entre los documentos del proyecto SCTIS.
"""
import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Paleta CORPOELEC
AZUL_OSCURO = RGBColor(0x00, 0x3C, 0x71)
AZUL_MEDIO = RGBColor(0x00, 0x50, 0x99)
GRIS_TEXTO = RGBColor(0x33, 0x33, 0x33)
GRIS_CLARO = RGBColor(0x55, 0x55, 0x55)
GRIS_MUY_CLARO = RGBColor(0x99, 0x99, 0x99)
NARANJA = RGBColor(0xCC, 0x66, 0x00)
VERDE = RGBColor(0x1A, 0x6E, 0x1A)
ROJO = RGBColor(0xCC, 0x00, 0x00)

SOMBRA_AZUL = "003C71"
SOMBRA_AZUL_SUAVE = "E8F0FE"
SOMBRA_CELESTE = "E8F4FD"
SOMBRA_VERDE = "E8F5E9"
SOMBRA_GRIS = "F0F4F8"
SOMBRA_GRIS_FILA = "F5F8FC"


def crear_documento():
    """Crea un Document con el estilo base del proyecto."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = GRIS_TEXTO

    t = doc.styles['Title']
    t.font.size = Pt(22)
    t.font.bold = True
    t.font.color.rgb = AZUL_OSCURO

    for nivel, (size, color) in [(1, (16, AZUL_OSCURO)), (2, (13, AZUL_MEDIO)), (3, (11.5, AZUL_MEDIO))]:
        h = doc.styles[f'Heading {nivel}']
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = color
        fmt = h.paragraph_format
        fmt.space_before = Pt(18 if nivel == 1 else 14 if nivel == 2 else 10)
        fmt.space_after = Pt(8 if nivel == 1 else 6 if nivel == 2 else 4)

    return doc


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith('shd'):
            tcPr.remove(child)
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    
    # OOXML schema sequence in tcPr: tcW, gridSpan, hMerge, vMerge, tcBorders, shd, noWrap, tcMar, vAlign
    vAlign = tcPr.find(qn('w:vAlign'))
    tcMar = tcPr.find(qn('w:tcMar'))
    noWrap = tcPr.find(qn('w:noWrap'))
    ref = vAlign or tcMar or noWrap
    if ref is not None:
        ref.addprevious(shading_elm)
    else:
        tcPr.append(shading_elm)


def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith('tcBorders'):
            tcPr.remove(child)
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "003C71")}"/>'
        )
        tcBorders.append(element)
        
    shd = tcPr.find(qn('w:shd'))
    vAlign = tcPr.find(qn('w:vAlign'))
    tcMar = tcPr.find(qn('w:tcMar'))
    noWrap = tcPr.find(qn('w:noWrap'))
    ref = shd or vAlign or tcMar or noWrap
    if ref is not None:
        ref.addprevious(tcBorders)
    else:
        tcPr.append(tcBorders)


def add_flow_table(doc, steps, header_color=SOMBRA_AZUL):
    """Tabla tipo flujograma con pasos en secuencia."""
    n = len(steps)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (title, desc) in enumerate(steps):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p_title = cell.paragraphs[0]
        run_t = p_title.add_run(title)
        run_t.bold = True
        run_t.font.size = Pt(9)
        run_t.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
        p_desc = cell.add_paragraph()
        run_d = p_desc.add_run(desc)
        run_d.font.size = Pt(7.5)
        run_d.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if n > 1:
        arrow_row = table.add_row()
        for i in range(n):
            cell = arrow_row.cells[i]
            if i < n - 1:
                run = cell.paragraphs[0].add_run("→")
                run.font.size = Pt(18)
                run.bold = True
                run.font.color.rgb = AZUL_OSCURO
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                cell.text = ""
    doc.add_paragraph("")


def add_branching_flow(doc, title, steps_left, steps_right, left_label="Opción A", right_label="Opción B"):
    """Flujo de dos ramas (decisión)."""
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = AZUL_MEDIO

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    left_cell = table.rows[0].cells[0]
    left_cell.text = ""
    p = left_cell.paragraphs[0]
    run_l = p.add_run(left_label)
    run_l.bold = True
    run_l.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for step in steps_left:
        sp = left_cell.add_paragraph()
        run_s = sp.add_run(f"• {step}")
        run_s.font.size = Pt(8)

    mid_cell = table.rows[0].cells[1]
    mid_cell.text = ""
    mp = mid_cell.paragraphs[0]
    mr = mp.add_run("⬇\n¿Decisión?")
    mr.bold = True
    mr.font.size = Pt(9)
    mr.font.color.rgb = NARANJA
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    right_cell = table.rows[0].cells[2]
    right_cell.text = ""
    p2 = right_cell.paragraphs[0]
    run_r = p2.add_run(right_label)
    run_r.bold = True
    run_r.font.size = Pt(9)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for step in steps_right:
        sp2 = right_cell.add_paragraph()
        run_s2 = sp2.add_run(f"• {step}")
        run_s2.font.size = Pt(8)

    for cell in [left_cell, mid_cell, right_cell]:
        set_cell_shading(cell, SOMBRA_GRIS)

    doc.add_paragraph("")


def add_info_box(doc, text, label="LENGUAJE CLARO", color=SOMBRA_CELESTE):
    """Caja informativa destacada."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    p_label = cell.paragraphs[0]
    r_label = p_label.add_run(f"  {label}")
    r_label.bold = True
    r_label.font.size = Pt(9)
    r_label.font.color.rgb = AZUL_OSCURO
    p_text = cell.add_paragraph()
    r_text = p_text.add_run(f"  {text}")
    r_text.font.size = Pt(10)
    r_text.font.italic = True
    r_text.font.color.rgb = GRIS_TEXTO
    set_cell_shading(cell, color)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    doc.add_paragraph("")


def add_tech_note(doc, text, label="NOTA TÉCNICA"):
    """Caja de nota técnica."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    p_label = cell.paragraphs[0]
    r_label = p_label.add_run(f"  {label}")
    r_label.bold = True
    r_label.font.size = Pt(9)
    r_label.font.color.rgb = VERDE
    p_text = cell.add_paragraph()
    r_text = p_text.add_run(f"  {text}")
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = GRIS_TEXTO
    set_cell_shading(cell, SOMBRA_VERDE)
    doc.add_paragraph("")


def add_decision_table(doc, steps, headers=("Paso", "Descripción", "Siguiente")):
    """Tabla de pasos con flechas."""
    table = doc.add_table(rows=len(steps) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, (step, desc, arrow) in enumerate(steps):
        row = table.rows[idx + 1]
        row.cells[0].text = step
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].text = desc
        row.cells[2].text = arrow
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph("")


def add_simple_table(doc, data, headers, col_widths=None, font_size=9):
    """Tabla genérica con encabezado azul y filas alternadas."""
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, SOMBRA_AZUL)
    for idx, row_data in enumerate(data):
        row = table.rows[idx + 1]
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            p = cell.paragraphs[0]
            for r in p.runs:
                r.font.size = Pt(font_size - 0.5)
        if idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, SOMBRA_GRIS_FILA)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    doc.add_paragraph("")
    return table


def add_portada(doc, titulo, subtitulo, sistema="SCTIS v1.0", fecha="Agosto 2026",
                dirigido=("Ing. Adrián Correa — Gerente de Gestión de Planificación de Distribución",
                          "Ing. Catherina Favio — Responsable del Proceso de Interrupciones"),
                autor="Yván Cipirán — Ingeniería de Productos de IA / IA Aplicada al SEN"):
    """Portada estándar del proyecto (sin logo, texto MPPEE/CORPOELEC)."""
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run("MPPEE  •  CORPOELEC")
    r_logo.font.size = Pt(14)
    r_logo.font.color.rgb = AZUL_OSCURO
    r_logo.bold = True

    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_org = p_org.add_run("Ministerio del Poder Popular para la Energía Eléctrica\nCorporación Eléctrica Nacional")
    r_org.font.size = Pt(11)
    r_org.font.color.rgb = GRIS_CLARO

    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_line = p_line.add_run("_" * 50)
    r_line.font.color.rgb = AZUL_OSCURO
    r_line.font.size = Pt(8)

    doc.add_paragraph("")

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_title.add_run(titulo)
    r_t.bold = True
    r_t.font.size = Pt(22)
    r_t.font.color.rgb = AZUL_OSCURO
    r_t.font.name = 'Calibri'

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(subtitulo + "\n" + sistema)
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = AZUL_MEDIO

    doc.add_paragraph("")
    doc.add_paragraph("")

    info_data = [
        ("Unidad Emisora:", "Servicios de Automatización de Procesos de Distribución"),
        ("Dirigido a:", dirigido[0]),
        ("", dirigido[1]),
        ("Gerencia:", "Gerencia General de Distribución\nGerencia de Gestión de Planificación de Distribución\nGrupo de Trabajo Seguimiento y Control"),
        ("Empresa:", "MPPEE / CORPOELEC"),
        ("Desarrollado por:", autor),
        ("Versión:", "1.0"),
        ("Fecha:", fecha),
    ]
    info_table = doc.add_table(rows=len(info_data), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info_data):
        cell_l = info_table.rows[i].cells[0]
        cell_v = info_table.rows[i].cells[1]
        cell_l.text = ""
        cell_v.text = ""
        r_l = cell_l.paragraphs[0].add_run(label)
        r_l.bold = True
        r_l.font.size = Pt(10)
        r_v = cell_v.paragraphs[0].add_run(value)
        r_v.font.size = Pt(10)
        if label:
            set_cell_shading(cell_l, SOMBRA_AZUL_SUAVE)
    for row in info_table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)

    doc.add_page_break()


def add_control_documento(doc, doc_codigo, titulo, version="1.0", fecha="Agosto 2026"):
    """Sección de control de documentos (histórico de versiones)."""
    doc.add_heading("Control de Documento", level=1)
    doc.add_paragraph("")
    # Tabla de metadatos
    meta = doc.add_table(rows=0, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Código", doc_codigo),
        ("Título", titulo),
        ("Versión", version),
        ("Fecha de emisión", fecha),
        ("Sistema", "SCTIS v1.0 — Sistema de Seguimiento y Control de Tiras de Interrupción"),
        ("Base de datos", "ggpd_se_cto_v1 (PostgreSQL / Supabase)"),
        ("Elaborado por", "Yván Cipirán — Ingeniería de Productos de IA / IA Aplicada al SEN"),
        ("Revisado por", "Ing. Catherina Favio — Responsable del Proceso de Interrupciones"),
        ("Aprobado por", "Ing. Adrián Correa — Gerente de Gestión de Planificación de Distribución"),
        ("Clasificación", "Uso interno CORPOELEC"),
    ]
    for label, value in meta_data:
        row = meta.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
        set_cell_shading(row.cells[0], SOMBRA_GRIS)

    doc.add_paragraph("")
    doc.add_heading("Historial de Versiones", level=2)
    add_simple_table(doc, [
        ("1.0", fecha, "Emisión inicial", "Yván Cipirán"),
    ], ["Versión", "Fecha", "Descripción", "Elaborado por"], col_widths=[Cm(2), Cm(3), Cm(8), Cm(3)])
    doc.add_page_break()


def pie_de_documento(doc, texto_extra=""):
    """Pie de documento con fecha de generación."""
    doc.add_paragraph("")
    p_note = doc.add_paragraph()
    extra = ("\n" + texto_extra) if texto_extra else ""
    r_note = p_note.add_run(
        "— Fin del documento —\n"
        "Documento generado el " + datetime.datetime.now().strftime("%d/%m/%Y %H:%M") +
        " por el Sistema de Automatización de Procesos de Distribución." + extra
    )
    r_note.font.size = Pt(8)
    r_note.font.color.rgb = GRIS_MUY_CLARO
    r_note.font.italic = True
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER


def agregar_glosario(doc, items):
    """Tabla de glosario término/definición."""
    table = doc.add_table(rows=len(items) + 1, cols=2)
    table.style = 'Table Grid'
    for i, h in enumerate(["Término", "Definición"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    for idx, (term, defn) in enumerate(items):
        row = table.rows[idx + 1]
        row.cells[0].text = term
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = defn
        for cell in row.cells:
            for p in cell.paragraphs:
                p.style.font.size = Pt(10)
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(12)
    doc.add_paragraph("")


def guardar(doc, ruta):
    doc.save(ruta)
    print(f"Documento generado: {ruta}")
