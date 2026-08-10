from docx import Document
import os
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ─── Page Setup ───
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ─── Styles ───
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Title style
title_style = doc.styles['Title']
title_font = title_style.font
title_font.size = Pt(22)
title_font.bold = True
title_font.color.rgb = RGBColor(0x00, 0x3C, 0x71)

# Heading 1
h1 = doc.styles['Heading 1']
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x00, 0x3C, 0x71)
h1_fmt = h1.paragraph_format
h1_fmt.space_before = Pt(18)
h1_fmt.space_after = Pt(8)

# Heading 2
h2 = doc.styles['Heading 2']
h2.font.size = Pt(13)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x00, 0x50, 0x99)
h2_fmt = h2.paragraph_format
h2_fmt.space_before = Pt(14)
h2_fmt.space_after = Pt(6)

# Heading 3
h3 = doc.styles['Heading 3']
h3.font.size = Pt(11.5)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x00, 0x50, 0x99)
h3_fmt = h3.paragraph_format
h3_fmt.space_before = Pt(10)
h3_fmt.space_after = Pt(4)

# ─── Helper Functions ───

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith('shd'):
            tcPr.remove(child)
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    
    # OOXML schema sequence in tcPr: tcW, gridSpan, hMerge, vMerge, tcBorders, shd, noWrap, tcMar, vAlign
    vAlign = tcPr.find(qn('w:vAlign'))
    tcMar = tcPr.find(qn('w:tcMar'))
    noWrap = tcPr.find(qn('w:noWrap'))
    ref = vAlign or tcMar or noWrap
    if ref is not None:
        ref.addprevious(shading_elm)
    else:
        tcPr.append(shading_elm)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith('tcBorders'):
            tcPr.remove(child)
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" '
            f'w:space="0" '
            f'w:color="{val.get("color", "003C71")}"/>'
        )
        tcBorders.append(element)
        
    shd = tcPr.find(qn('w:shd'))
    vAlign = tcPr.find(qn('w:vAlign'))
    tcMar = tcPr.find(qn('w:tcMar'))
    noWrap = tcPr.find(qn('w:noWrap'))
    ref = shd or vAlign or tcMar or noWrap
    if ref is not None:
        ref.addprevious(tcBorders)
    else:
        tcPr.append(tcBorders)

def add_flow_table(doc, steps, header_color="003C71"):
    """Add a flowchart-style table with steps in sequence."""
    n = len(steps)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (title, desc) in enumerate(steps):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p_title = cell.paragraphs[0]
        run_t = p_title.add_run(title)
        run_t.bold = True
        run_t.font.size = Pt(9)
        run_t.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
        # Description below
        p_desc = cell.add_paragraph()
        run_d = p_desc.add_run(desc)
        run_d.font.size = Pt(7.5)
        run_d.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Arrow after cell (except last)
        if i < n - 1:
            arrow_cell = table.rows[0].cells[i]  # can't easily add between cells
    # Add arrow row below
    if n > 1:
        arrow_row = table.add_row()
        for i in range(n):
            cell = arrow_row.cells[i]
            if i < n - 1:
                run = cell.paragraphs[0].add_run("→")
                run.font.size = Pt(18)
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x3C, 0x71)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                cell.text = ""
    doc.add_paragraph("")

def add_branching_flow(doc, title, steps_left, steps_right, left_label="Opción A", right_label="Opción B"):
    """Two-branch flow table."""
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x50, 0x99)
    
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Left steps
    left_cell = table.rows[0].cells[0]
    left_cell.text = ""
    p = left_cell.paragraphs[0]
    run_l = p.add_run(left_label)
    run_l.bold = True
    run_l.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for step in steps_left:
        sp = left_cell.add_paragraph()
        run_s = sp.add_run(f"• {step}")
        run_s.font.size = Pt(8)
    
    # Middle: arrow
    mid_cell = table.rows[0].cells[1]
    mid_cell.text = ""
    mp = mid_cell.paragraphs[0]
    mr = mp.add_run("⬇\n¿Decisión?")
    mr.bold = True
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Right steps
    right_cell = table.rows[0].cells[2]
    right_cell.text = ""
    p2 = right_cell.paragraphs[0]
    run_r = p2.add_run(right_label)
    run_r.bold = True
    run_r.font.size = Pt(9)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for step in steps_right:
        sp2 = right_cell.add_paragraph()
        run_s2 = sp2.add_run(f"• {step}")
        run_s2.font.size = Pt(8)
    
    # Set borders
    for cell in [left_cell, mid_cell, right_cell]:
        set_cell_shading(cell, "F0F4F8")
    
    doc.add_paragraph("")

def add_info_box(doc, text, label="LENGUAJE CLARO", color="E8F4FD"):
    """Add an info/callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    p_label = cell.paragraphs[0]
    r_label = p_label.add_run(f"  {label}")
    r_label.bold = True
    r_label.font.size = Pt(9)
    r_label.font.color.rgb = RGBColor(0x00, 0x3C, 0x71)
    
    p_text = cell.add_paragraph()
    r_text = p_text.add_run(f"  {text}")
    r_text.font.size = Pt(10)
    r_text.font.italic = True
    r_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    set_cell_shading(cell, color)
    cell_par = cell.paragraphs[0]
    cell_par.paragraph_format.space_before = Pt(3)
    cell_par.paragraph_format.space_after = Pt(3)
    doc.add_paragraph("")

def add_tech_note(doc, text, label="NOTA TÉCNICA"):
    """Add technical note box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    p_label = cell.paragraphs[0]
    r_label = p_label.add_run(f"  {label}")
    r_label.bold = True
    r_label.font.size = Pt(9)
    r_label.font.color.rgb = RGBColor(0x1A, 0x6E, 0x1A)
    
    p_text = cell.add_paragraph()
    r_text = p_text.add_run(f"  {text}")
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    set_cell_shading(cell, "E8F5E9")
    doc.add_paragraph("")

def add_decision_table(doc, steps):
    """Table with step number, description, arrows."""
    table = doc.add_table(rows=len(steps)+1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    headers = ["Paso", "Descripción", "Siguiente"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for idx, (step, desc, arrow) in enumerate(steps):
        row = table.rows[idx + 1]
        row.cells[0].text = step
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].text = desc
        row.cells[2].text = arrow
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    
    doc.add_paragraph("")


# ═══════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════

doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")

# Shield / logo text
p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_logo = p_logo.add_run("MPPEE  •  CORPOELEC")
r_logo.font.size = Pt(14)
r_logo.font.color.rgb = RGBColor(0x00, 0x3C, 0x71)
r_logo.bold = True

p_org = doc.add_paragraph()
p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_org = p_org.add_run("Ministerio del Poder Popular para la Energía Eléctrica\nCorporación Eléctrica Nacional")
r_org.font.size = Pt(11)
r_org.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_line = p_line.add_run("─" * 60)
r_line.font.color.rgb = RGBColor(0x00, 0x3C, 0x71)
r_line.font.size = Pt(8)

doc.add_paragraph("")

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_t = p_title.add_run("DOCUMENTO DE ARQUITECTURA\nY FLUJOS DE PROCESO")
r_t.bold = True
r_t.font.size = Pt(22)
r_t.font.color.rgb = RGBColor(0x00, 0x3C, 0x71)
r_t.font.name = 'Calibri'

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run("Sistema de Gestión de Tiras de Interrupción Eléctrica\nSCTIS v1.0")
r_sub.font.size = Pt(14)
r_sub.font.color.rgb = RGBColor(0x00, 0x50, 0x99)

doc.add_paragraph("")
doc.add_paragraph("")

# Information table
info_table = doc.add_table(rows=7, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

info_data = [
    ("Unidad Emisora:", "Servicios de Automatización de Procesos de Distribución"),
    ("Dirigido a:", "Ing. Adrián Correa — Gerente de Gestión de Planificación de Distribución"),
    ("", "Ing. Catherina Favio — Responsable del Proceso de Interrupciones"),
    ("Gerencia:", "Gerencia General de Distribución\nGerencia de Gestión de Planificación de Distribución\nGrupo de Trabajo Seguimiento y Control"),
    ("Empresa:", "MPPEE / CORPOELEC"),
    ("Versión:", "1.0"),
    ("Fecha:", "Julio 2026"),
]

for i, (label, value) in enumerate(info_data):
    cell_l = info_table.rows[i].cells[0]
    cell_v = info_table.rows[i].cells[1]
    cell_l.text = ""
    cell_v.text = ""
    r_l = cell_l.paragraphs[0].add_run(label)
    r_l.bold = True
    r_l.font.size = Pt(10)
    r_v = cell_v.paragraphs[0].add_run(value)
    r_v.font.size = Pt(10)
    if label:
        set_cell_shading(cell_l, "E8F0FE")

# Set column widths
for row in info_table.rows:
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(11)

doc.add_page_break()

# ═══════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════

doc.add_heading("Índice", level=1)
toc_items = [
    "1. Diagrama de Contexto General",
    "2. Proceso de Carga Masiva desde Excel",
    "   2.1. Subir Archivo",
    "   2.2. Validación de Estructura",
    "   2.3. Detección de Activos Inconsistentes",
    "   2.4. Homologación de Causas",
    "   2.5. Previsualización y Confirmación",
    "3. Proceso de Carga Manual (Formulario Web)",
    "4. Proceso de Control de Acceso y Seguridad",
    "5. Proceso de Control de Calidad ISO 8000",
    "6. Esquema General de la Base de Datos",
    "7. Glosario Rápido",
    "8. Tareas Pendientes para Definir con el Funcional",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if not item.startswith("   "):
        p.runs[0].bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════
# 1. DIAGRAMA DE CONTEXTO GENERAL
# ═══════════════════════════════════════════════

doc.add_heading("1. Diagrama de Contexto General", level=1)

add_flow_table(doc, [
    ("USUARIO\nOPERADOR", "Sube Excel\nLlena formulario"),
    ("SCTIS\nAPP WEB", "Valida estructura\nResuelve activos\nHomologa causas"),
    ("BASE DE\nDATOS", "PostgreSQL\nRLS por estado\nAuditoría"),
    ("USUARIO\nADMIN", "Gestiona\nusuarios"),
])

doc.add_paragraph(
    "El sistema recibe datos de interrupción desde dos vías: un archivo Excel "
    "(importación masiva de múltiples registros) o un formulario web (carga manual "
    "de una sola tira). En ambos casos, la aplicación valida automáticamente que los "
    "datos cumplan con la estructura ISO 8000, resuelve inconsistencias de activos "
    "(subestaciones y circuitos no registrados), homologa las causas contra el catálogo "
    "oficial de 16 causas, y solo entonces persiste en la base de datos. Cada usuario "
    "ve exclusivamente los registros de su estado asignado gracias al filtro RLS "
    "(Row Level Security) a nivel de PostgreSQL."
)

# ═══════════════════════════════════════════════
# 2. PROCESO DE CARGA MASIVA DESDE EXCEL
# ═══════════════════════════════════════════════

doc.add_heading("2. Proceso de Carga Masiva desde Excel", level=1)

doc.add_paragraph(
    "Este es el flujo principal del sistema. El planificador de cada estado descarga "
    "el Formato Establecido (plantilla Excel de 22 columnas), consolida sus 8 reportes "
    "mensuales en una sola hoja, y la sube al sistema. El flujo completo consta de "
    "7 pasos secuenciales, desde la subida del archivo hasta la inserción definitiva "
    "en la base de datos."
)

# Overview flow
add_flow_table(doc, [
    ("PASO 1", "Subir\nArchivo"),
    ("PASO 2", "Validar\nEstructura"),
    ("PASO 3", "Detectar\nActivos"),
    ("PASO 4", "Resolver\nActivos"),
    ("PASO 5", "Homologar\nCausas"),
    ("PASO 6", "Preview y\nConfirmar"),
    ("PASO 7", "Insertar\nen BD"),
])

# 2.1
doc.add_heading("2.1. Paso 1 — Subir Archivo", level=2)

doc.add_paragraph(
    "El usuario accede a la opción 'Importar' en el menú principal, selecciona el "
    "archivo Excel (.xlsx) desde su equipo y hace clic en 'Subir'. El sistema recibe "
    "el archivo en el servidor y lo prepara para la validación estructural."
)

add_info_box(doc,
    "¿Qué espera el sistema? Un archivo Excel con el Formato Establecido de 22 columnas. "
    "No se aceptan formatos libres. Cada planificador debe consolidar sus reportes en "
    "esta plantilla antes de subir. Si el archivo no coincide, el sistema lo rechaza "
    "con un mensaje claro indicando qué columna falta o está incorrecta."
)

# 2.2
doc.add_heading("2.2. Paso 2 — Validación de Estructura", level=2)

add_flow_table(doc, [
    ("LEER\nARCHIVO", "Abre Excel\ny extrae datos"),
    ("VERIFICAR\nCOLUMNAS", "22 columnas,\norden exacto"),
    ("VALIDAR\nTIPOS", "Fechas, números,\ntexto"),
    ("¿ERROR?", "Rechazo con\ninstructivo"),
    ("¿OK?", "Pasa a\nPaso 3"),
])

doc.add_paragraph(
    "El sistema examina el archivo en tres niveles: (1) cantidad de columnas y nombres "
    "esperados, (2) tipos de datos de cada columna, (3) valores obligatorios no nulos. "
    "Si alguna columna clave falta, está desplazada, o contiene tipos de datos "
    "incorrectos (ej: una fecha almacenada como texto), el sistema devuelve un error "
    "422 con un instructivo detallado de la corrección requerida."
)

add_tech_note(doc,
    "Técnicamente: se valida contra COLUMNAS_FORMULARIO (22 columnas) y "
    "HEADERS_FORMATO_ESPERADOS. La verificación es posición-dependente: el orden de "
    "las columnas importa tanto como su nombre. Esto evita desplazamientos accidentales "
    "que corromperían los datos aguas abajo."
)

# 2.3
doc.add_heading("2.3. Paso 3 — Detección de Activos Inconsistentes", level=2)

add_flow_table(doc, [
    ("EXTRAER\nACTIVOS", "Lee subestaciones\ny circuitos\ndel Excel"),
    ("CONSULTAR\nINVENTARIO", "Cruza contra\ncommon.assets\nen BD"),
    ("¿EXISTE\nSUBESTACIÓN?", "Busca por\ncódigo y estado"),
    ("¿EXISTE\nCIRCUITO?", "Verifica parentesco\ncon subestación"),
    ("INCONSISTENCIA\nDETECTADA", "Notifica al\nusuario"),
])

doc.add_paragraph(
    "El sistema cruza cada par (subestación, circuito) del archivo Excel contra el "
    "inventario oficial de activos de CORPOELEC registrado en common.assets. Si "
    "aparece una subestación que no está registrada, o un circuito que no pertenece "
    "a esa subestación, el sistema lo detecta y ofrece al usuario dos opciones."
)

add_branching_flow(doc, "Resolución por el usuario:",
    ["Crear nuevo activo en BD", "Se registra en common.assets", "Queda disponible para futuras cargas"],
    ["Mapear a activo existente", "Se asigna manualmente", "Equivalente a corregir el nombre"],
    left_label="Opción A: Crear nuevo",
    right_label="Opción B: Mapear a existente"
)

add_info_box(doc,
    "La resolución de activos es en cascada: primero se resuelve la subestación, y una vez "
    "resuelta, se validan los circuitos que cuelgan de ella. No se puede resolver un "
    "circuito si su subestación padre sigue siendo inconsistente.",
    label="REGLA DE NEGOCIO"
)

# 2.4
doc.add_heading("2.4. Paso 4 — Homologación de Causas", level=2)

doc.add_paragraph(
    "Cada estado reporta las causas de interrupción con su propio criterio y nomenclatura. "
    "El sistema homologa todas esas variantes a las 16 causas oficiales definidas por "
    "la Gerencia. Actualmente hay 96 registros de causas cargados y homologados al 100%."
)

add_decision_table(doc, [
    ("1", "Extraer causa del archivo, normalizar texto", "→ 2"),
    ("2", "Buscar coincidencia exacta en catálogo de 96 variantes", "→ 3a o 3b"),
    ("3a", "Coincidencia exacta → Homologación automática", "→ 4"),
    ("3b", "Sin coincidencia exacta → El sistema sugiere la causa oficial más cercana", "→ 3c"),
    ("3c", "Usuario confirma la sugerencia o selecciona manualmente de las 16 causas", "→ 4"),
    ("4", "Causa homologada y asociada al registro", "→ Paso 5"),
])

add_tech_note(doc,
    "El catálogo de homologación reside en sctis.causa. Cada variante de causa reportada "
    "por los estados tiene un mapeo a una de las 16 causas oficiales. El matching es "
    "case-insensitive y usa normalización Unicode para manejar acentos y caracteres especiales."
)

# 2.5
doc.add_heading("2.5. Paso 5 — Previsualización y Confirmación", level=2)

add_flow_table(doc, [
    ("GENERAR\nPREVIEW", "N° registros\nActivos creados\nCausas homologadas"),
    ("ALERTAS DE\nCALIDAD", "Scoring ISO\nDuplicados\nAdvertencias"),
    ("USUARIO\n¿CONFIRMA?", "Sí → Insertar\nNo → Descartar"),
    ("INSERTAR\nEN BD", "Commit transaccional\nRegistro de auditoría"),
])

doc.add_paragraph(
    "Antes de escribir un solo registro en la base de datos, el sistema presenta un "
    "resumen al usuario: cantidad de tiras a insertar, activos nuevos creados, causas "
    "homologadas, y cualquier alerta de calidad (duplicados detectados, campos con "
    "bajo scoring). El usuario decide si confirma la importación o la descarta. "
    "No se persiste nada hasta la confirmación explícita. Si confirma, toda la "
    "operación se ejecuta en una transacción: o se inserta todo o no se inserta nada."
)

add_info_box(doc,
    "Caso de uso: si el usuario detecta en el preview que la duración de varios "
    "registros parece incorrecta, puede cancelar la importación, corregir el Excel "
    "y subirlo de nuevo sin dejar datos huérfanos en la BD. Es una red de seguridad.",
    label="LENGUAJE CLARO"
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 3. PROCESO DE CARGA MANUAL
# ═══════════════════════════════════════════════

doc.add_heading("3. Proceso de Carga Manual (Formulario Web)", level=1)

doc.add_paragraph(
    "Para cargas de una sola tira o correcciones puntuales, el sistema dispone de un "
    "formulario web interactivo con selecciones en cascada (cascading selects) que "
    "guían al usuario y previenen errores de referencias."
)

add_decision_table(doc, [
    ("1", "Usuario accede a /nuevo", "→ 2"),
    ("2", "Estado: bloqueado según perfil del usuario (solo puede cargar su estado)", "→ 3"),
    ("3", "Selecciona Subestación: filtrada por el estado (cascading)", "→ 4"),
    ("4", "Selecciona Circuito: filtrado por la subestación seleccionada", "→ 5"),
    ("5", "Selecciona Despachador: filtrado por el estado", "→ 6"),
    ("6", "Llena campos: fecha, inicio, duración, carga, TTI, causa, etc.", "→ 7"),
    ("7", "Click 'Guardar' → Validaciones ISO", "→ 8"),
    ("8", "¿Todo OK? → Inserta en BD. ¿Error? → Muestra mensaje al usuario", "fin"),
])

add_info_box(doc,
    "Piénselo como un asistente: selecciona el estado y el sistema le muestra solo "
    "las subestaciones de ese estado; selecciona la subestación y le muestra solo "
    "sus circuitos. No puede equivocarse de referencia porque las opciones disponibles "
    "se ajustan automáticamente.",
    label="LENGUAJE CLARO"
)

add_tech_note(doc,
    "Las listas en cascada se implementan con endpoints REST: "
    "GET /api/subestaciones?estado=X, GET /api/circuitos?subestacion_id=Y, "
    "GET /api/despachadores?estado=X. El frontend consume estas APIs con fetch "
    "asíncrono y actualiza los <select> sin recargar la página."
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 4. CONTROL DE ACCESO Y SEGURIDAD
# ═══════════════════════════════════════════════

doc.add_heading("4. Proceso de Control de Acceso y Seguridad", level=1)

doc.add_paragraph(
    "El sistema implementa seguridad en tres capas: autenticación (verificar quién es), "
    "autorización (verificar qué puede hacer), y aislamiento de datos (RLS a nivel BD)."
)

# Authentication sub-flow
doc.add_heading("Autenticación", level=2)

add_decision_table(doc, [
    ("1", "Usuario ingresa usuario y contraseña", "→ 2"),
    ("2", "¿Usuario existe en BD?", "No → Error 'Credenciales inválidas'"),
    ("3", "¿Usuario está activo?", "No → Error 'Usuario bloqueado/desactivado'"),
    ("4", "¿Contraseña correcta?", "No → Incrementa contador de intentos fallidos"),
    ("4a", "¿5 intentos fallidos?", "Sí → Bloqueo de 30 minutos"),
    ("5", "Autenticación exitosa → Crea sesión persistente", "→ Registro en audit.access_log"),
])

add_tech_note(doc,
    "El bloqueo por intentos fallidos se implementa con una política de 5 intentos "
    "en 30 minutos. El contador se resetea al iniciar sesión exitosamente o al "
    "cumplirse la ventana de tiempo. Los intentos se registran en audit.access_log "
    "con timestamp, IP y user-agent."
)

# Authorization sub-flow
doc.add_heading("Autorización", level=2)

add_flow_table(doc, [
    ("SESIÓN\nINICIADA", "Usuario autenticado\nRol: admin o editor"),
    ("¿ROL\nADMIN?", "Ve TODOS los\nestados"),
    ("¿ROL\nEDITOR?", "Ve SOLO su estado\nasignado"),
    ("RLS EN\nBD", "Filtro obligatorio\na nivel PostgreSQL"),
])

doc.add_paragraph(
    "Cada usuario tiene un perfil en user_profiles con estado_codigo asignado y un "
    "rol (admin, editor). El admin puede ver, editar y gestionar todos los estados. "
    "El editor solo ve y opera sobre los registros de su estado asignado. Este filtro "
    "no es solo visual: está implementado con Row Level Security (RLS) en PostgreSQL, "
    "lo que significa que aunque un usuario acceda directamente a la BD, las políticas "
    "de seguridad lo impiden a nivel de fila."
)

add_info_box(doc,
    "Punto clave para auditoría: cada acción queda registrada en audit.access_log con "
    "usuario, timestamp, acción realizada, y datos involucrados. La trazabilidad es "
    "completa e irreversible.",
    label="PARA LA AUDITORÍA"
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 5. CONTROL DE CALIDAD ISO 8000
# ═══════════════════════════════════════════════

doc.add_heading("5. Proceso de Control de Calidad ISO 8000", level=1)

doc.add_paragraph(
    "El sistema implementa controles de calidad de datos basados en el estándar ISO 8000, "
    "adaptado al contexto de tiras de interrupción eléctrica."
)

add_flow_table(doc, [
    ("CAMPOS\nOBLIGATORIOS", "Verifica que no\nhaya nulos en\ncolumnas clave"),
    ("TIPOS DE\nDATOS", "Fechas, números,\nformatos correctos"),
    ("INTEGRIDAD\nREFERENCIAL", "FKs: estado, subestación,\ncircuito, causa,\ndespachador"),
    ("SCORING\nDE CALIDAD", "Puntaje 0-100\npor registro"),
    ("DETECCIÓN\nDUPLICADOS", "Misma fecha +\ncircuito + hora"),
])

doc.add_paragraph(
    "Cada registro recibe un puntaje de calidad (scoring) basado en completitud, "
    "consistencia y precisión de sus campos. Los registros con baja calidad no se "
    "bloquean (porque muchas veces la información disponible es la única que existe), "
    "pero quedan marcados y registrados en audit.data_issues para revisión posterior. "
    "Los duplicados se detectan por combinación de fecha + circuito + hora de inicio "
    "y se notifican al usuario antes de la inserción."
)

add_tech_note(doc,
    "Estructura de audit.data_issues: id_tira FK, tipo_issue (duplicado, scoring_bajo, "
    "campo_faltante), severidad (alta, media, baja), descripcion, creado_en. "
    "Esto permite generar reportes de calidad por estado, por período, o por tipo de "
    "incidencia."
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 6. ESQUEMA DE BASE DE DATOS
# ═══════════════════════════════════════════════

doc.add_heading("6. Esquema General de la Base de Datos", level=1)

doc.add_paragraph(
    "El sistema utiliza PostgreSQL 17 con tres esquemas principales:"
)

# Tables as simple description
schemas = [
    ("common", "Datos maestros compartidos", [
        "states — catálogo de estados (24 registros)",
        "assets — subestaciones y circuitos (inventario CORPOELEC)",
    ]),
    ("sctis", "Datos del sistema de interrupciones", [
        "tira_interrupcion — registros de interrupción (tabla principal)",
        "causa — catálogo de 16 causas oficiales + 96 variantes homologadas",
        "despachador — despachadores asociados a cada estado",
        "user_profiles — perfil, rol y estado asignado por usuario",
        "user_roles — roles del sistema (admin, editor)",
    ]),
    ("audit", "Auditoría y calidad", [
        "access_log — inicio de sesión, intentos fallidos, actividad",
        "change_log — modificaciones a registros existentes",
        "data_issues — incidencias de calidad ISO 8000",
    ]),
]

for schema_name, schema_desc, tables in schemas:
    h = doc.add_heading(f"Esquema {schema_name}", level=2)
    p = doc.add_paragraph(schema_desc)
    for t in tables:
        bp = doc.add_paragraph(t, style='List Bullet')

doc.add_paragraph(
    "Las relaciones entre tablas se gestionan mediante claves foráneas (FK) y políticas "
    "RLS. El esquema común (common.assets) incluye un campo parent_id que permite la "
    "jerarquía subestación → circuitos. Todas las tablas transaccionales tienen "
    "timestamp de creación y modificación."
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 7. GLOSARIO
# ═══════════════════════════════════════════════

doc.add_heading("7. Glosario Rápido", level=1)

glossary = [
    ("TTI", "Tiempo Total de Interrupción, expresado en MWh. Representa la energía no suministrada durante el evento."),
    ("RLS", "Row Level Security. Mecanismo de PostgreSQL que restringe qué filas puede ver/modificar cada usuario según su perfil."),
    ("Cascading Select", "Lista desplegable que se actualiza automáticamente según la selección anterior (ej: al elegir subestación, se filtran sus circuitos)."),
    ("Scoring ISO", "Puntaje de calidad del registro (0-100) basado en completitud, consistencia y precisión de los datos."),
    ("Homologación", "Proceso de mapear una causa reportada por un estado a una de las 16 causas oficiales del catálogo."),
    ("Activo", "Subestación o circuito registrado en el inventario CORPOELEC (tabla common.assets)."),
    ("Formato Establecido", "Plantilla Excel oficial de 22 columnas que deben usar todos los estados para reportar sus tiras de interrupción."),
    ("FK (Foreign Key)", "Clave foránea que garantiza integridad referencial entre dos tablas."),
]

table = doc.add_table(rows=len(glossary)+1, cols=2)
table.style = 'Table Grid'

# Header
for i, h in enumerate(["Término", "Definición"]):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for idx, (term, defn) in enumerate(glossary):
    row = table.rows[idx + 1]
    row.cells[0].text = term
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[1].text = defn
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(12)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 8. TAREAS PENDIENTES
# ═══════════════════════════════════════════════

doc.add_heading("8. Tareas Pendientes para Definir con el Funcional", level=1)

doc.add_paragraph(
    "A continuación se listan los puntos que requieren decisión por parte del equipo "
    "funcional para continuar con el desarrollo y ajuste del sistema. Cada tarea "
    "incluye la descripción del problema, su dependencia y prioridad sugerida."
)

pending_tasks = [
    ("1", "Definir columnas de fecha/hora del Formato Establecido",
     "Determinar si el formulario de 22 columnas debe incluir HORA INICIO (G), "
     "FECHA FIN (I) y HORA FIN (J) además de FECHA INICIO (H). Actualmente las "
     "columnas G, I y V están en blanco. Los estados Zulia, Miranda y La Guaira "
     "ya manejan fecha/hora separadas en sus reportes.",
     "Funcional (Ing. Catherina)", "Alta"),
    ("2", "Definir contenido de la columna V",
     "Actualmente en blanco en el Formato Establecido. Posibles usos: tensión "
     "nominal, tipo de falta, o código de identificación del activo.",
     "Funcional", "Media"),
    ("3", "Validar catálogo de 16 causas homologadas",
     "Confirmar que las 96 variantes de causas cargadas y mapeadas a las 16 "
     "oficiales cubren todos los escenarios que reportan los estados. "
     "Identificar si faltan causas o sobran.",
     "Jefe de Unidad + Ing. Catherina", "Alta"),
    ("4", "Definir alcance de auditoría de transformaciones",
     "Decidir si se auditan todos los campos calculados (TTI, duración, MVAmín) "
     "o solo los de mayor impacto. Propuesta: crear tabla audit.transformaciones "
     "para registrar el valor original vs. valor final por campo, permitiendo "
     "trazabilidad completa ante auditorías externas.",
     "Ing. Adrián + Ing. Catherina", "Media"),
    ("5", "Estrategia de migración a Google AI Studio",
     "Definir hoja de ruta para migrar la aplicación Flask actual a Google AI "
     "Studio cuando sea aprobada por la Gerencia.",
     "Ing. Adrián", "Baja"),
    ("6", "Plan de despliegue por estado",
     "Cronograma de capacitación, prueba y activación para cada uno de los "
     "14 estados operadores que usarán el sistema.",
     "Jefe de Unidad", "Media"),
    ("7", "Revisión de calidad de datos históricos",
     "Evaluar si los datos ya cargados en la BD requieren ajustes menores "
     "(errores de tipeo, causas mal homologadas) antes de poner el sistema "
     "en producción oficial.",
     "Ing. Catherina", "Media"),
]

p_table = doc.add_table(rows=len(pending_tasks)+1, cols=5)
p_table.style = 'Table Grid'
p_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
p_headers = ["#", "Tarea", "Descripción", "Depende de", "Prioridad"]
for i, h in enumerate(p_headers):
    cell = p_table.rows[0].cells[i]
    cell.text = h
    p = cell.paragraphs[0]
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, "003C71")

for idx, (num, task, desc, dep, prio) in enumerate(pending_tasks):
    row = p_table.rows[idx + 1]
    data = [num, task, desc, dep, prio]
    for i, val in enumerate(data):
        cell = row.cells[i]
        cell.text = val
        p = cell.paragraphs[0]
        run = p.runs[0]
        run.font.size = Pt(8.5)
        if i == 4:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if val == "Alta":
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                run.bold = True
            elif val == "Media":
                run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)
            else:
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif i == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Alternating row colors
for idx in range(1, len(pending_tasks)+1):
    if idx % 2 == 0:
        for cell in p_table.rows[idx].cells:
            set_cell_shading(cell, "F5F8FC")

# Column widths
col_widths = [Cm(0.8), Cm(3.5), Cm(7.5), Cm(3.5), Cm(1.5)]
for row in p_table.rows:
    for i in range(5):
        row.cells[i].width = col_widths[i]

# ─── Footer note ───
doc.add_paragraph("")
p_note = doc.add_paragraph()
r_note = p_note.add_run(
    "— Fin del documento —\n"
    "Documento generado el " + datetime.datetime.now().strftime("%d/%m/%Y %H:%M") + " "
    "por el Sistema de Automatización de Procesos de Distribución."
)
r_note.font.size = Pt(8)
r_note.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
r_note.font.italic = True
p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ─── Save ───
def docx_to_html_doc(doc, title, out_path):
    parts = []
    parts.append('<!DOCTYPE html>')
    parts.append('<html xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:w="urn:schemas-microsoft-com:office:word" xmlns="http://www.w3.org/TR/REC-html40">')
    parts.append('<head><meta charset="utf-8"><title>' + title + '</title>')
    parts.append('<style>')
    parts.append('body { font-family: Arial, sans-serif; font-size: 11pt; color: #1e293b; line-height: 1.6; margin: 40px; }')
    parts.append('h1 { color: #003C71; font-size: 20pt; margin-top: 24pt; margin-bottom: 12pt; border-bottom: 2px solid #003C71; padding-bottom: 6px; }')
    parts.append('h2 { color: #003C71; font-size: 15pt; margin-top: 18pt; margin-bottom: 8pt; }')
    parts.append('h3 { color: #003C71; font-size: 12pt; margin-top: 14pt; margin-bottom: 6pt; }')
    parts.append('p { margin-bottom: 10pt; text-align: justify; }')
    parts.append('table { border-collapse: collapse; width: 100%; margin: 14pt 0; }')
    parts.append('th, td { border: 1px solid #cbd5e1; padding: 8pt 10pt; text-align: left; vertical-align: top; font-size: 10pt; }')
    parts.append('th { background-color: #003C71; color: #ffffff; font-weight: bold; }')
    parts.append('tr:nth-child(even) td { background-color: #f8fafc; }')
    parts.append('</style></head><body>')

    for element in doc.element.body:
        tag = element.tag.split('}')[-1]
        if tag == 'p':
            runs = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            text = ''.join([r.text for r in runs if r.text])
            if text.strip():
                p_style = element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
                val = p_style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if p_style is not None else ''
                if 'Heading1' in val or 'Heading 1' in val or 'Title' in val:
                    parts.append(f'<h1>{text.strip()}</h1>')
                elif 'Heading2' in val or 'Heading 2' in val:
                    parts.append(f'<h2>{text.strip()}</h2>')
                elif 'Heading3' in val or 'Heading 3' in val:
                    parts.append(f'<h3>{text.strip()}</h3>')
                else:
                    parts.append(f'<p>{text.strip()}</p>')
        elif tag == 'tbl':
            parts.append('<table>')
            rows = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
            for r_idx, r in enumerate(rows):
                parts.append('<tr>')
                cells = r.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
                tag_name = 'th' if r_idx == 0 else 'td'
                for c in cells:
                    txts = [t.text for t in c.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t') if t.text]
                    c_text = ' '.join(txts).strip()
                    parts.append(f'<{tag_name}>{c_text}</{tag_name}>')
                parts.append('</tr>')
            parts.append('</table>')

    parts.append('</body></html>')
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(parts))


for out_dir in ['/app/applet/interrupciones/docs', '/app/applet/docs']:
    os.makedirs(out_dir, exist_ok=True)
    out_file = os.path.join(out_dir, "FLUJO_PROCESOS_SCTIS_v1.docx")
    out_doc_file = os.path.join(out_dir, "FLUJO_PROCESOS_SCTIS_v1.doc")
    doc.save(out_file)
    docx_to_html_doc(doc, "FLUJO DE PROCESOS SCTIS", out_doc_file)
    print(f"Documento generado: {out_file}")
    print(f"Documento .doc (Google Drive) generado: {out_doc_file}")

