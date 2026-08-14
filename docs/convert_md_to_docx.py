import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D3D3D3')
        tblBorders.append(border)
    border = OxmlElement('w:insideV')
    border.set(qn('w:val'), 'none')
    tblBorders.append(border)
    tblPr.append(tblBorders)

def md_to_docx(md_path, docx_path, doc_path=None):
    doc = docx.Document()
    
    # Page setup (Margins 1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                code_text = "".join(code_lines)
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = table.cell(0, 0)
                set_cell_background(cell, 'F4F6F8')
                cell.width = Inches(6.5)
                
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Pt(6)
                
                run = p.add_run(code_text.strip())
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                
                doc.add_paragraph() # Spacing
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue
            
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables (Markdown format)
        if '|' in stripped and ('---' in stripped or i + 1 < len(lines) and '|' in lines[i+1]):
            # Collect table lines
            table_lines = []
            while i < len(lines) and '|' in lines[i].strip():
                table_lines.append(lines[i].strip())
                i += 1
            
            # Process table
            rows_data = []
            for tline in table_lines:
                if set(tline.replace('|', '').replace(':', '').replace('-', '').strip()) == set():
                    continue
                cols = [c.strip() for c in tline.split('|')[1:-1]]
                if cols:
                    rows_data.append(cols)
                    
            if rows_data:
                num_rows = len(rows_data)
                num_cols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table)
                
                for r_idx, row in enumerate(rows_data):
                    for c_idx, val in enumerate(row):
                        if c_idx < num_cols:
                            cell = table.cell(r_idx, c_idx)
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(3)
                            p.paragraph_format.space_after = Pt(3)
                            
                            # Clean markdown bold/italics
                            clean_val = re.sub(r'\*\*(.*?)\*\*', r'\1', val)
                            clean_val = re.sub(r'\*(.*?)\*', r'\1', clean_val)
                            
                            run = p.add_run(clean_val)
                            run.font.name = 'Calibri'
                            
                            if r_idx == 0:
                                set_cell_background(cell, '003366') # CORPOELEC Navy
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.font.size = Pt(10)
                            else:
                                if r_idx % 2 == 1:
                                    set_cell_background(cell, 'FFFFFF')
                                else:
                                    set_cell_background(cell, 'F8FAFC')
                                run.font.size = Pt(9.5)
                                run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                                if '**' in val:
                                    run.font.bold = True
                
                doc.add_paragraph() # Spacing after table
            continue

        # Headings
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[2:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x66) # CORPOELEC Navy
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[3:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x44, 0x88)
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped[4:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(12.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        elif stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped[5:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        # Horizontal Rule
        elif stripped in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("_________________________________________________________________________________")
            run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
            run.font.size = Pt(8)
        # Bullet list
        elif stripped.startswith('* ') or stripped.startswith('- ') or stripped.startswith('• '):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            # Format bold/inline text
            parts = re.split(r'(\*\*.*?\*\*)', bullet_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        # Regular paragraph
        elif stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            # Format bold/inline text
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        i += 1
        
    doc.save(docx_path)
    print(f"Saved: {docx_path}")
    
    # Save a .doc copy as well if requested
    if doc_path:
        doc.save(doc_path)
        print(f"Saved: {doc_path}")

if __name__ == '__main__':
    base_dir = "/home/skidrowkodex/Documentos/Repositorio_Maestro/docs"
    files = [
        "NAC_2026_GGPD_AUDITORIA_ESTATUS_CUATRO_APLICACIONES_V01",
        "NAC_2026_GGPD_AUDITORIA_TECNICA_GOBERNANZA_BD_V01",
        "NAC_2026_GGPD_AUDITORIA_FUNCIONAL_GOBIERNO_DATOS_V01",
        "NAC_2026_GGPD_AUDITORIA_TECNICA_CONFORMIDAD_ISO_COBIT_V01",
        "NAC_2026_GGPD_AUDITORIA_FUNCIONAL_MATRIZ_GOBERNANZA_V01",
    ]
    qa_files = [
        "NAC_2026_GGPD_INVENTARIO_ARQUITECTURA_RUTAS_DESPLIEGUE_V01",
        "NAC_2026_GGPD_RESUMEN_EJECUTIVO_DESPLIEGUE_USUARIOS_QA_V01",
        "NAC_2026_GGPD_MATRIZ_25_CUENTAS_VISOR_ESTADAL_SIGI_V01"
    ]
    
    for f in files:
        md_file = os.path.join(base_dir, f + ".md")
        docx_file = os.path.join(base_dir, f + ".docx")
        doc_file = os.path.join(base_dir, f + ".doc")
        if os.path.exists(md_file):
            md_to_docx(md_file, docx_file, doc_file)

    qa_dir = os.path.join(base_dir, "despliegues_qa")
    for f in qa_files:
        md_file = os.path.join(qa_dir, f + ".md")
        docx_file = os.path.join(qa_dir, f + ".docx")
        doc_file = os.path.join(qa_dir, f + ".doc")
        if os.path.exists(md_file):
            md_to_docx(md_file, docx_file, doc_file)
